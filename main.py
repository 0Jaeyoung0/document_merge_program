import src.file_io.file_io as file_io
import src.convert_document.convert_to_pdf as convert_to_pdf
import src.merge_document.merge_pdf as merge_pdf
import os
'index'
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

'gui'
from tkinter import *
import tkinter
from tkinter.messagebox import *
from tkinter import messagebox
import os



def btn_load_click(event):
    global input_files
    input_files = file_selector.open_files()
    for file in input_files:
        file_basename = os.path.basename(file)  # 파일 이름만 추출
        listbox.insert(END, file_basename)  # 리스트박스에 항목 추가
    return input_files

def btn_delete_click(event):
    selected_items = listbox.curselection()  # 선택된 항목의 인덱스를 가져옴
    if askokcancel("Delete", f"Are you sure you want to delete it?"):
        print("yes")
        # 선택된 항목을 뒤에서부터 삭제하여 인덱스 변화를 방지
        for i in reversed(selected_items):
            listbox.delete(i)  # 리스트박스에서 항목 삭제
            del input_files[i]  # 전역 리스트에서 항목 삭제

    else:
        print("no")

def btn_merge_click(event, input_files):
    print("merge")

    converted_files = []
    for index, input_file in enumerate(input_files, start=1):
        file_path, file_name = os.path.split(input_file)
        converted_file = os.path.join(file_path, f'temp{index}.pdf')

        to_pdf_converter.convert_to_pdf(input_file=input_file, output_file=converted_file)
        # if(input_file.endswith('.doc')) or input_file.endswith('.docx'):
        #    to_pdf_converter.extract_page(converted_file, input_file)

        converted_files.append(converted_file)

    output_file = file_selector.save_file()

    pdf_merger.merge_pdf(input_files=converted_files, output_file=output_file)

    for converted_file in converted_files:
        os.remove(converted_file)
    listbox.delete(0, END)
    input_files = []
    tkinter.messagebox.showinfo("Complete","Merged is complete.")


if __name__ == '__main__':

    file_selector = file_io.FileIO()
    to_pdf_converter = convert_to_pdf.ToPdfConverter()
    pdf_merger = merge_pdf.PdfMerger()

    root = Tk()
    root.title("N&M")  # 창 제목 설정
    root.resizable(width=False, height=False)  # 창 크기 조절 불가능하도록 설정
    root.geometry("400x200+100+100")

    listbox = Listbox(root, width=20)  # 리스트박스 생성
    listbox.pack()
    # listbox.grid(row=0, column=0, columnspan=3, sticky='ew')  # 리스트박스 위치 조정
    listbox.pack(side="top", fill="x")

    # 버튼 생성 및 위치 조정
    btn_load = Button(root, text='파일 로드', width=10)
    btn_delete = Button(root, text='파일 삭제', width=10)
    btn_merge = Button(root, text='파일 병합', width=10)
    btn_load.pack(side="left")
    btn_delete.pack(side="left")
    btn_merge.pack(side="right")

    # 이벤트 바인딩
    btn_load.bind('<Button-1>', btn_load_click)
    btn_delete.bind('<Button-1>', btn_delete_click)
    btn_merge.bind('<Button-1>', lambda event: btn_merge_click(event, input_files))  # 이벤트 바인딩 시 input_files 전달

    root.mainloop()

'index'
'''
    document = Document()


    title = document.add_heading('INDEX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    num_files = 3
    files_pages = (2,3,1)
    files_description = ('first file', 'second file', 'third file')


    document.sections[0].left_margin = Cm(2.6)
    document.sections[0].right_margin = Cm(2)
    document.sections[0].top_margin = Cm(2)
    document.sections[0].bottom_margin = Cm(2)

    table = document.add_table(rows=num_files+1,cols=4)
    table.style = document.styles['Table Grid']
    heading_cells = table.rows[0].cells
    heading_text = ('No', 'DESCRIPTION', 'PAGE','REMARKS')
    heading_width = (1.44, 10.25 ,2.27, 3.41)

    #heading cell
    for i, (text, width) in enumerate(zip(heading_text, heading_width)):
        heading_cells[i].text = text
        heading_cells[i].width = Cm(width)
        heading_cells[i].paragraphs[0].style.font.name = '맑은 고딕'
        if i > 0:
            heading_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #content cell
    row = table.rows[1]
    row.cells[0].text = 'row = 1, column = 0'
    row.cells[1].text = 'row = 1, column = 1'
    total_page_num = 3
    for i in range(1, num_files+1):
        row = table.rows[i]
        row.cells[0].text = f'{i}'
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = files_description[i-1]
        row.cells[2].text = f'P{total_page_num}~{total_page_num+files_pages[i-1]+2}'
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_page_num = total_page_num+files_pages[i-1]
    document.save('test.docx')
'''