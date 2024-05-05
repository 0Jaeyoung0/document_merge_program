from docx import Document
from docx.shared import Inches

class WordHandler(object):
    def __init__(self):
        pass

    @staticmethod
    def create_cover_page(doc_name, dept_name, person_name, background_image):
        document = Document()

        # 배경 이미지 설정
        WordHandler.set_background_image(document, background_image)

        # 문서 제목 추가
        document.add_heading(doc_name, 0)

        # 빈 줄 추가
        document.add_paragraph('')

        # 담당 부서 추가
        department = document.add_paragraph()
        department_run = department.add_run('담당 부서: ')
        department_run.bold = True
        department.add_run(dept_name)

        # 담당자 추가
        person = document.add_paragraph()
        person_run = person.add_run('담당자: ')
        person_run.bold = True
        person.add_run(person_name)

        # 페이지 여백 설정
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        return document

    @staticmethod
    def set_background_image(document, image_path):



if __name__ == '__main__':
    doc = WordHandler.create_cover_page('프로젝트 계획서', '기획팀', '홍길동', '../../sample_data/image_sample/samsung.png')
    doc.save('cover_page.docx')